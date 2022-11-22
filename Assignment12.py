"""

1. In what modes should the PdfFileReader() and PdfFileWriter() File objects will be opened?

These files will be opened in binary mode, read binary (rb) for PdfFileReader() and write binary
(wb) PdfWriter()

2. From a PdfFileReader object, how do you get a Page object for page 5?

Calling getPage(4) will return a page object for page 5 since page 0 is the first page.


3. What PdfFileReader variable stores the number of pages in the PDF document?
The numPages variable stores an integer of the number of pages in the PdfFileReader object.

4. If a PdfFileReader object’s PDF is encrypted with the password swordfish, what must you do before
you can obtain Page objects from it?

call decrypt('swordfish')


5. What methods do you use to rotate a page?

The rotateClockwise() and rotateCounterClockwise() methods.


6. What is the difference between a Run object and a Paragraph object?

 A document contains multiple paragraphs. A paragraph begins on a new line and contains multiple 
 runs. Runs are contiguous groups of characters within the paragraphs.

7. How do you obtain a list of Paragraph objects for a Document object that’s stored in a variable
named doc?

By using doc.paragraphs

8. What type of object has bold, underline, italic, strike, and outline variables?

A Run object has bold, underline,italic, strike and outline variables.

9. What is the difference between False, True, and None for the bold variable?

True is the attribute is always enables, no matter what other styles are applied to the run.
False is the attribute is always disabled.
None is defaults to whatever the run's style is set to.


10. How do you create a Document object for a new Word document?

By calling docx.Documnet() function.

11. How do you add a paragraph with the text 'Hello, there!' to a Document object stored in a
variable named doc?

import docx
doc=docx.Document()
doc.add_paragraph('Hello there!')
doc.save('hellothere.docx')


12. What integers represent the levels of headings available in Word documents?
 integer 0 to 4

"""